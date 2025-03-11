import docx
import uuid
import json
from loguru import logger
from app.core.config import (
    MODEL_PATH,
    TABLE_NAME,
    TEXT_SPLITTER_CONFIG,
    EMBEDDING_CONFIG,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import LanceDB  # noqa: F401
import lancedb  # type: ignore
from app.core.db import init_database


class WordLoader:
    def __init__(self, model_path=MODEL_PATH):
        self.embedding = HuggingFaceEmbeddings(
            model_name=str(model_path), **EMBEDDING_CONFIG
        )
        self.text_splitter = RecursiveCharacterTextSplitter(**TEXT_SPLITTER_CONFIG)
        self.connection = init_database()
        self.file_path = None

    def _read_word_file(self, file_path):
        """读取Word文档内容，并提取文本和段落信息"""
        doc = docx.Document(file_path)

        # 提取所有段落文本
        paragraphs_text = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # 跳过空段落
                paragraphs_text.append(
                    {"text": para.text, "metadata": {"paragraph": i + 1}}
                )

        # 提取表格内容
        for i, table in enumerate(doc.tables):
            table_text = f"表格 {i+1}:\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                table_text += row_text + "\n"

            paragraphs_text.append({"text": table_text, "metadata": {"table": i + 1}})

        self.file_path = file_path

        return paragraphs_text

    def _split_text(self, documents):
        """分割文本并保留元数据"""
        return self.text_splitter.split_documents(
            [
                Document(page_content=doc["text"], metadata=doc["metadata"])
                for doc in documents
            ]
        )

    def _vector_exists(self, vector_store, text):
        results = vector_store.similarity_search_with_score(
            text, k=1, score_threshold=0.99
        )
        return len(results) > 0

    def _write_to_lancedb(self, texts, source):
        connection = lancedb.connect(self.connection.uri)
        table = connection.open_table(TABLE_NAME)

        for text in texts:
            # 使用embedding模型获取向量
            vector = self.embedding.embed_query(text.page_content)

            # 准备元数据
            metadata = {}
            if "paragraph" in text.metadata:
                metadata["paragraph"] = text.metadata["paragraph"]
            elif "table" in text.metadata:
                metadata["table"] = text.metadata["table"]
            metadata["file_path"] = self.file_path

            # 准备数据
            data = {
                "id": str(uuid.uuid4()),
                "vector": vector,
                "text": text.page_content,
                "source": source,
                "metadata": json.dumps(metadata),
            }

            # 添加到表中
            table.add([data])
        logger.info(f"文件{self.file_path}已写入数据库")

    @logger.catch
    def process_file(self, file_path):
        word_content = self._read_word_file(file_path)
        split_docs = self._split_text(word_content)
        self._write_to_lancedb(split_docs, file_path)


if __name__ == "__main__":
    loader = WordLoader()
    loader.process_file(r"D:\Test_Data\sample.docx")
