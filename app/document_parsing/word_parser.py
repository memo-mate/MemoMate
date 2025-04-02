import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings
from app.core.log_adapter import logger
from app.document_parsing.common import (
    CHINESE_NUMERALS,
    COMMON_TITLE_KEYWORDS,
    TITLE_PATTERNS,
    DocumentContentType,
    DocumentSection,
    ParserConfig,
)


class WordParser:
    """Word文档解析器"""

    def __init__(self, config: ParserConfig = None):
        self.config = config or ParserConfig()

    def __call__(
        self, file_path: str | Path, *args: Any, **kwargs: Any
    ) -> List[LangchainDocument]:
        """处理文档并返回分块后的文档列表"""
        return self.chunk(file_path)

    def chunk(self, file_path: str | Path) -> List[LangchainDocument]:
        """将文档分割成多个块"""
        file_path = Path(file_path) if not isinstance(file_path, Path) else file_path

        if file_path.suffix.lower() != ".docx":
            logger.error("不支持的文件类型", file_path=str(file_path), expected="docx")
            raise ValueError(f"不支持的文件类型: {file_path.suffix}, 仅支持docx文件")

        try:
            if self.config.strategy == ParserConfig.STRATEGY_STRUCTURE:
                return self._structure_split(file_path)
            else:
                return self._basic_split(file_path)
        except Exception as e:
            raise ValueError(f"无法解析Word文件: {str(e)}")

    def _basic_split(self, file_path: Path) -> List[LangchainDocument]:
        """基本的文档分块"""
        doc = Document(file_path)

        text, metadata = self._extract_text_and_metadata(doc, file_path)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", ";", ":", "  ", " ", ""],
        )

        chunks = text_splitter.create_documents([text], [metadata])

        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
            chunk.metadata["chunk_count"] = len(chunks)

            lines = chunk.page_content.split("\n")
            for line in lines[:3]:
                line = line.strip()
                if line and len(line) < 100 and not line.endswith("."):
                    chunk.metadata["possible_title"] = line
                    break

        logger.info("基本分块完成", file_path=str(file_path), chunk_count=len(chunks))
        return chunks

    def _structure_split(self, file_path: Path) -> List[LangchainDocument]:
        """根据文档结构进行分块"""

        doc = Document(file_path)

        _, base_metadata = self._extract_text_and_metadata(doc, file_path)

        sections = self._analyze_document_structure(doc)

        if not sections:
            logger.warning("未检测到文档结构，回退到基本分块", file_path=str(file_path))
            return self._basic_split(file_path)

        documents = []

        for section in sections:
            if section.level in self.config.heading_levels:
                metadata = base_metadata.copy()
                metadata.update(
                    {
                        "title": section.title,
                        "level": section.level,
                        "content_type": DocumentContentType.TEXT.value,
                    }
                )

                section_text = f"# {section.title}\n\n{section.content}"

                if len(section_text) <= settings.CHUNK_SIZE * 1.2:
                    documents.append(
                        LangchainDocument(page_content=section_text, metadata=metadata)
                    )
                else:
                    try:
                        text_splitter = RecursiveCharacterTextSplitter(
                            chunk_size=settings.CHUNK_SIZE,
                            chunk_overlap=settings.CHUNK_OVERLAP,
                            length_function=len,
                            separators=[
                                "\n\n",
                                "\n",
                                ". ",
                                "! ",
                                "? ",
                                ";",
                                ":",
                                "  ",
                                " ",
                                "",
                            ],
                        )
                        sub_chunks = text_splitter.create_documents(
                            [section_text], [metadata]
                        )

                        if sub_chunks and not sub_chunks[0].page_content.startswith(
                            f"# {section.title}"
                        ):
                            title_text = f"# {section.title}\n\n"
                            sub_chunks[0].page_content = (
                                title_text + sub_chunks[0].page_content
                            )

                        for chunk in sub_chunks:
                            chunk.metadata["title"] = section.title
                            chunk.metadata["level"] = section.level

                        documents.extend(sub_chunks)
                    except Exception as e:
                        logger.warning(
                            "章节分块失败，作为整块处理",
                            section_title=section.title,
                            error=str(e),
                        )
                        documents.append(
                            LangchainDocument(
                                page_content=section_text, metadata=metadata
                            )
                        )

        if not documents:
            logger.warning(
                "未生成任何结构化文档块，回退到基本分块", file_path=str(file_path)
            )
            return self._basic_split(file_path)

        for i, doc in enumerate(documents):
            doc.metadata["chunk_index"] = i
            doc.metadata["chunk_count"] = len(documents)

        logger.info(
            "结构化分块完成", file_path=str(file_path), chunk_count=len(documents)
        )
        return documents

    def _extract_text_and_metadata(
        self, doc, file_path: Path
    ) -> Tuple[str, Dict[str, Any]]:
        """提取Word文本内容和元数据"""
        metadata = {
            "source": str(file_path),
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower().lstrip("."),
        }

        if hasattr(doc, "core_properties"):
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.created:
                metadata["created"] = doc.core_properties.created.isoformat()
            if doc.core_properties.modified:
                metadata["modified"] = doc.core_properties.modified.isoformat()

        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                text_content.append("\n".join(table_text))

        full_text = "\n\n".join(text_content)

        return full_text, metadata

    def _analyze_document_structure(self, doc) -> List[DocumentSection]:
        """分析文档结构以识别标题和章节"""
        sections = []
        heading_styles = self._identify_heading_styles(doc)
        current_section = None

        current_sections = {}

        paragraphs = []
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            style_name = para.style.name if hasattr(para.style, "name") else ""
            content_type = DocumentContentType.TEXT.value
            title_level = 0

            if style_name in heading_styles:
                content_type = DocumentContentType.TITLE.value
                title_level = heading_styles[style_name]

            if content_type != DocumentContentType.TITLE.value:
                title_info = self._parse_title_format(para.text)
                if title_info:
                    number, title_text, level = title_info
                    content_type = DocumentContentType.TITLE.value
                    title_level = level

            paragraphs.append(
                {
                    "text": para.text,
                    "content_type": content_type,
                    "title_level": title_level
                    if content_type == DocumentContentType.TITLE.value
                    else 0,
                }
            )

        tables = []
        for i, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))

            if table_text:
                tables.append(
                    {
                        "text": "\n".join(table_text),
                        "content_type": DocumentContentType.TABLE.value,
                        "table_index": i,
                    }
                )

        elements = paragraphs + tables
        elements.sort(
            key=lambda x: x.get("table_index", 0)
            if x["content_type"] == DocumentContentType.TABLE.value
            else 0
        )

        for i, element in enumerate(elements):
            if element["content_type"] == DocumentContentType.TITLE.value:
                level = element["title_level"]

                new_section = DocumentSection(
                    title=element["text"],
                    level=level,
                    content="",
                    metadata={"elements": []},
                )

                sections.append(new_section)

                current_sections[level] = new_section

                higher_levels = [l for l in current_sections.keys() if l > level]
                for l in higher_levels:
                    current_sections.pop(l, None)

                current_section = new_section
            elif current_section:
                if current_section.content:
                    current_section.content += f"\n\n{element['text']}"
                else:
                    current_section.content = element["text"]
                current_section.metadata["elements"].append(element)
            else:
                default_section = DocumentSection(
                    title="文档开始",
                    level=0,
                    content=element["text"],
                    metadata={"elements": [element]},
                )
                sections.append(default_section)
                current_sections[0] = default_section
                current_section = default_section

        return sections

    def _identify_heading_styles(self, doc) -> Dict[str, int]:
        """识别文档中使用的标题样式"""
        heading_styles = {}

        # 常见的标题样式名称模式
        heading_patterns = [
            (r"^副标题\s*字符$", lambda x: 2),  # 精确匹配副标题 字符
            (r"[Hh]eading\s*(\d+)", lambda x: int(x)),  # Heading 1, heading2, ...
            (r"标题\s*(\d+)", lambda x: int(x)),  # 标题 1, 标题2, ...
            (r"(\d+)\s*级标题", lambda x: int(x)),  # 2 级标题, 3级标题 等
            (r"[Tt]itle\s*(\d+)", lambda x: int(x)),  # Title 1, title2, ...
            (r"[Tt]itle", lambda x: 1),  # Title
            (r"标题", lambda x: 1),  # 标题
            (r"副标题", lambda x: 2),  # 副标题
        ]

        for style in doc.styles:
            if not hasattr(style, "name"):
                continue

            style_name = style.name

            for pattern, level_func in heading_patterns:
                match = re.search(pattern, style_name)
                if match:
                    try:
                        level = (
                            level_func(match.group(1))
                            if match.groups()
                            else level_func(None)
                        )
                        heading_styles[style_name] = level
                        break
                    except (IndexError, ValueError):
                        heading_styles[style_name] = 1

        return heading_styles

    def _parse_title_format(self, text: str) -> Optional[Tuple[str, str, int]]:
        """解析标题格式，返回(编号, 标题文本, 层级)或None"""
        for pattern, level in TITLE_PATTERNS:
            match = re.match(pattern, text)
            if match:
                number = match.group(1)
                title_text = match.group(2).strip()

                if len(title_text.split()) > 20 or any(
                    p in title_text for p in ["。", "；", ";", "."]
                ):
                    continue

                if re.search(r"[一二三四五六七八九十百]", number):
                    try:
                        if number == "十":
                            num = 10
                        elif "十" in number:
                            parts = number.split("十")
                            if parts[0]:
                                num = CHINESE_NUMERALS[parts[0]] * 10
                            else:
                                num = 10
                            if parts[1]:
                                num += CHINESE_NUMERALS[parts[1]]
                        elif "百" in number:
                            parts = number.split("百")
                            if parts[0]:
                                num = CHINESE_NUMERALS[parts[0]] * 100
                            else:
                                num = 100
                            if parts[1]:
                                if "十" in parts[1]:
                                    tens_parts = parts[1].split("十")
                                    if tens_parts[0]:
                                        num += CHINESE_NUMERALS[tens_parts[0]] * 10
                                    else:
                                        num += 10
                                    if tens_parts[1]:
                                        num += CHINESE_NUMERALS[tens_parts[1]]
                                else:
                                    num += CHINESE_NUMERALS[parts[1]]
                        else:
                            num = CHINESE_NUMERALS[number]
                        number = str(num)
                    except (KeyError, ValueError):
                        continue

                return number, title_text, level

        if len(text) < 40 and not text.endswith(
            ("。", "，", "；", "：", ".", ",", ";", ":")
        ):
            if text in COMMON_TITLE_KEYWORDS:
                return "", text, 1

            for title in COMMON_TITLE_KEYWORDS:
                if text.startswith(title) and len(text) < len(title) + 10:
                    return "", text, 1

        return None


if __name__ == "__main__":
    parser = WordParser()
    docs = parser("C:/Users/Leo/Desktop/第二章 开发平台概览.docx")
